"""Generate DEAP Emotion Recognition technical report as .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Colour palette ───────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE    = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE  = RGBColor(0x44, 0x72, 0xC4)
GRAY        = RGBColor(0x59, 0x59, 0x59)
GREEN_FILL  = "E2F0D9"
RED_FILL    = "FFE0E0"
BLUE_FILL   = "DEEAF1"
HEADER_FILL = "1F3864"
TABLE_HEAD  = "D5E3F0"
YELLOW_FILL = "FFF2CC"


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, size=10, color=None, align="left", italic=False):
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # padding
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "80")
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tcPr.append(mar)
    set_cell_border(cell)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = DARK_BLUE if level == 1 else MED_BLUE
    run.font.size = Pt(16) if level == 1 else Pt(13)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=11, align="left", space_before=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == "justify" else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E75B6")
    pBdr.append(bot)
    pPr.append(pBdr)


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after  = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Распознавание эмоций\nна датасете DEAP")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(28)
run.font.color.rgb = DARK_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(4)
p2.paragraph_format.space_after  = Pt(4)
run2 = p2.add_run("Технический отчёт — NeuroBarometer Project")
run2.font.name = "Arial"
run2.font.size = Pt(14)
run2.font.color.rgb = MED_BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(2)
p3.paragraph_format.space_after  = Pt(30)
run3 = p3.add_run("Апрель 2026")
run3.font.name = "Arial"
run3.font.size = Pt(12)
run3.font.color.rgb = GRAY
run3.italic = True

add_rule(doc)

# ════════════════════════════════════════════════════════════════════════════
# 1. ПОСТАНОВКА ЗАДАЧИ
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Постановка задачи")
add_para(doc, (
    "Цель проекта — воспроизвести SOTA-результаты распознавания эмоций на датасете DEAP "
    "с использованием трёх модальностей: ЭЭГ (EEG), фотоплетизмограф (PPG) и "
    "кожно-гальваническая реакция (GSR/EDA). Пайплайн разрабатывается локально и служит "
    "основой для последующего внедрения в устройство NeuroBarometer."
), size=11)
add_para(doc, (
    "Задача: бинарная классификация Valence (позитивная / негативная эмоция) "
    "и Arousal (высокое / низкое возбуждение). Протоколы оценки: "
    "Subject-Dependent (SD, 8-fold CV) и Leave-One-Subject-Out (LOSO)."
), size=11)

# ════════════════════════════════════════════════════════════════════════════
# 2. ДАТАСЕТ
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Датасет DEAP")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.autofit = False
t.columns[0].width = Cm(5)
t.columns[1].width = Cm(11)
hdr = t.rows[0].cells
cell_text(hdr[0], "Параметр",   bold=True, size=10); set_cell_bg(hdr[0], TABLE_HEAD)
cell_text(hdr[1], "Значение",   bold=True, size=10); set_cell_bg(hdr[1], TABLE_HEAD)

dataset_rows = [
    ("Субъекты",     "32 человека (s01–s32)"),
    ("Trials",       "40 видеоклипов на субъект"),
    ("Каналы",       "40 всего: 0–31 ЭЭГ, 36 GSR, 38 PPG"),
    ("Частота",      "128 Гц, препроцессированные данные"),
    ("Длина trial",  "63 сек = 3 с базовая линия + 60 с стимул"),
    ("Метки",        "Valence, Arousal, Dominance, Liking [1–9]"),
    ("Бинаризация",  "Медиана субъекта (не фиксированный порог 5.0)"),
    ("Окна",         "60 окон × 1 сек = 2400 сэмплов / субъект"),
]
for k, v in dataset_rows:
    r = t.add_row().cells
    cell_text(r[0], k, size=10)
    cell_text(r[1], v, size=10)

# ════════════════════════════════════════════════════════════════════════════
# 3. ПРИЗНАКИ
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Признаковое пространство")
add_para(doc, (
    "Все признаки извлекаются в скользящих окнах 1 секунда. GroupKFold по trials "
    "предотвращает утечку данных между фолдами."
), size=11)

t2 = doc.add_table(rows=1, cols=3)
t2.style = "Table Grid"
t2.autofit = False
t2.columns[0].width = Cm(5)
t2.columns[1].width = Cm(1.5)
t2.columns[2].width = Cm(9.5)
for i, txt in enumerate(["Модальность", "Размер", "Описание"]):
    cell_text(t2.rows[0].cells[i], txt, bold=True, size=10)
    set_cell_bg(t2.rows[0].cells[i], TABLE_HEAD)

feat_rows = [
    ("EEG — Differential Entropy",     "160", "32 канала x 5 полос (delta, theta, alpha, beta, gamma)",    None),
    ("PPG — HRV признаки",             "10",  "SDNN, RMSSD, pNN50, LF, HF, LF/HF и др. (3 сегм. × 20 с)", None),
    ("GSR — EDA признаки",             "8",   "SCL mean/std, SCR пики, amplitude, slope (3 сегм. × 20 с)",  None),
    ("Консенсус по видео (NEW)",       "2",   "Средняя оценка каждого видео по остальным 31 субъектам",    GREEN_FILL),
    ("Позиция окна в trial",           "1",   "Нормализованная позиция 0→1 внутри 60-секундного trial",     None),
    ("Фронтальная асимметрия (FAA)",   "2",   "Alpha и theta асимметрия: правое − левое полушарие",         None),
    ("ИТОГО",                          "183", "Входная размерность LightMLP",                               YELLOW_FILL),
]
for name, sz, desc, fill in feat_rows:
    r = t2.add_row().cells
    cell_text(r[0], name, size=10, bold=(fill == YELLOW_FILL))
    cell_text(r[1], sz,   size=10, align="center", bold=(fill == YELLOW_FILL))
    cell_text(r[2], desc, size=10, bold=(fill == YELLOW_FILL))
    if fill:
        for c in r:
            set_cell_bg(c, fill)

# ════════════════════════════════════════════════════════════════════════════
# 4. АРХИТЕКТУРА
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Архитектура модели — LightMLP (~28K параметров)")
add_para(doc, (
    "Использована облегчённая MLP вместо Transformer-архитектур. Обоснование: "
    "35 обучающих trials × 60 окон = 2100 сэмплов на fold — "
    "недостаточно для сходимости тяжёлых моделей (660K+ параметров)."
), size=11)

t3 = doc.add_table(rows=1, cols=2)
t3.style = "Table Grid"
t3.autofit = False
t3.columns[0].width = Cm(6)
t3.columns[1].width = Cm(10)
for i, txt in enumerate(["Слой", "Параметры"]):
    cell_text(t3.rows[0].cells[i], txt, bold=True, size=10)
    set_cell_bg(t3.rows[0].cells[i], TABLE_HEAD)

arch_rows = [
    ("Input",                               "183 признака"),
    ("Dense → BN → GELU → Dropout(0.5)",   "183 → 256"),
    ("Dense → BN → GELU → Dropout(0.3)",   "256 → 128"),
    ("Valence head",                        "128 → 2 (Low / High)"),
    ("Arousal head",                        "128 → 2 (Low / High)"),
    ("Оптимизатор",                         "AdamW, lr=1e-3, CosineAnnealing, EarlyStop(15)"),
    ("GPU",                                 "NVIDIA RTX 4070 Ti SUPER"),
]
for k, v in arch_rows:
    r = t3.add_row().cells
    cell_text(r[0], k, size=10)
    cell_text(r[1], v, size=10)

# ════════════════════════════════════════════════════════════════════════════
# 5. ABLATION
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Ablation Study — вклад каждого признака")
add_para(doc, "Каждая гипотеза тестировалась накопительно. SD, 32 субъекта, majority vote.", size=11)

t4 = doc.add_table(rows=1, cols=5)
t4.style = "Table Grid"
t4.autofit = False
for w, h in zip([7, 2, 2, 2, 3], ["Конфигурация", "Val %", "Ar %", "Delta Val", "Статус"]):
    pass
widths = [Cm(7.2), Cm(1.8), Cm(1.8), Cm(2.0), Cm(3.2)]
for i, (col, txt) in enumerate(zip(t4.columns, ["Конфигурация", "Val %", "Ar %", "Delta Val", "Статус"])):
    col.width = widths[i]
    cell_text(t4.rows[0].cells[i], txt, bold=True, size=10)
    set_cell_bg(t4.rows[0].cells[i], TABLE_HEAD)

ablation = [
    ("Базовый LightMLP (DE + HRV + EDA)",             "66.6", "64.2", "—",      "Baseline",    None),
    ("+ H1: Групповой консенсус по видео",             "79.6", "70.9", "+13.0",  "Принято",     GREEN_FILL),
    ("+ H2: Baseline DE нормализация",                 "78.2", "69.2", "-1.4",   "Откат",       RED_FILL),
    ("+ H3: Позиция окна в trial",                     "79.9", "70.1", "+0.3",   "Нейтрально",  None),
    ("+ H4: Фронтальная FAA/theta асимметрия",         "79.7", "70.4", "-0.2",   "Нейтрально",  None),
    ("+ H5: Сегментированные PPG/GSR (3 × 20 с)",     "82.2", "71.8", "+2.3",   "Принято",     GREEN_FILL),
]
for cfg, val, ar, dv, status, fill in ablation:
    r = t4.add_row().cells
    for col, (txt, al) in enumerate(zip([cfg, val, ar, dv, status],
                                        ["left","center","center","center","center"])):
        cell_text(r[col], txt, size=10, align=al)
        if fill:
            set_cell_bg(r[col], fill)

# ════════════════════════════════════════════════════════════════════════════
# 6. РЕЗУЛЬТАТЫ
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Финальные результаты")

t5 = doc.add_table(rows=1, cols=5)
t5.style = "Table Grid"
t5.autofit = False
col_widths5 = [Cm(4.8), Cm(3.2), Cm(3.2), Cm(2.6), Cm(2.2)]
for i, (col, txt) in enumerate(zip(t5.columns, ["Протокол", "Valence Acc", "Arousal Acc", "Valence F1", "Arousal F1"])):
    col.width = col_widths5[i]
    cell_text(t5.rows[0].cells[i], txt, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(t5.rows[0].cells[i], HEADER_FILL)

results_rows = [
    ("SD (8-fold CV, 32 субъекта)",             "82.19 ± 8.31%",  "71.80 ± 11.15%", "82.24%",  "71.74%",  BLUE_FILL),
    ("LOSO (32 folds, 32 субъекта)",            "79.84 ± 8.84%",  "64.92 ± 10.54%", "79.78%",  "64.26%",  "EBF3E8"),
    ("Базовый MLP [SD]",                        "66.6%",          "64.2%",          "—",       "—",       None),
    ("EEGNet multimodal [SD]",                  "65.3%",          "63.0%",          "—",       "—",       None),
    ("SOTA статьи (SD, с data leakage)",        "88–99%",         "87–99%",         "—",       "—",       YELLOW_FILL),
]
for proto, val, ar, vf1, af1, fill in results_rows:
    r = t5.add_row().cells
    bold_row = fill in (BLUE_FILL, "EBF3E8")
    for col, (txt, al) in enumerate(zip([proto, val, ar, vf1, af1],
                                        ["left","center","center","center","center"])):
        cell_text(r[col], txt, size=10, align=al, bold=bold_row)
        if fill:
            set_cell_bg(r[col], fill)

# ════════════════════════════════════════════════════════════════════════════
# 7. ЧЕСТНОСТЬ ПРОТОКОЛА
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Корректность протокола оценки")
add_heading(doc, "Принятые меры против data leakage", level=2)
for b in [
    "GroupKFold по trials (не по окнам) — исключает пересечение между train и test",
    "StandardScaler.fit() только на train-части каждого fold",
    "Бинарный порог = медиана субъекта, не фиксированный 5.0",
    "Консенсус по видео вычисляется из ДРУГИХ субъектов, не из test-субъекта",
    "Majority vote по 60 окнам trial для финального предсказания",
]:
    add_bullet(doc, b)

add_heading(doc, "Почему SOTA статьи дают 90%+", level=2)
add_para(doc, "Большинство публикаций содержат неосознанную утечку данных:", size=11)
for b in [
    "Случайный split по окнам вместо по trials",
    "Нормализация на всём датасете до разделения",
    "Порог 5.0 создаёт 70-80% мажоритарный класс у части субъектов",
]:
    add_bullet(doc, b)

# ════════════════════════════════════════════════════════════════════════════
# 8. ДОРОЖНАЯ КАРТА
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Дорожная карта NeuroBarometer")
add_heading(doc, "Текущий статус", level=2)
for b in [
    "Воспроизводимый пайплайн на DEAP: SD 82.2% Valence / 71.8% Arousal",
    "LOSO 79.8% Valence / 64.9% Arousal — честная cross-subject оценка",
    "Полный код: deap_emotion/ (config, loader, features, models, trainer, main.py)",
]:
    add_bullet(doc, b)

add_heading(doc, "Следующие шаги при появлении собственных данных", level=2)
for b in [
    "Собрать 20–30 сессий с NeuroBarometer (EEG 32 кан. + PPG + GSR)",
    "Разметка: видео-стимулы с SAM шкалой (Valence / Arousal), как в DEAP",
    "Finetune LightMLP с собственными данными — ожидаемый результат 66–72%",
    "При 50+ субъектах — протестировать предобученные модели (LaBraM, EEGPT)",
    "Консенсус по видео применим при персонализации: накопленные данные нескольких пользователей",
]:
    add_bullet(doc, b)

add_rule(doc)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_foot.add_run("NeuroBarometer Project  |  2026  |  Конфиденциально")
run_f.font.name = "Arial"
run_f.font.size = Pt(9)
run_f.font.color.rgb = GRAY
run_f.italic = True

# ── Save ─────────────────────────────────────────────────────────────────────
out = "results/DEAP_Report_2026.docx"
doc.save(out)
print(f"Report saved: {out}")
